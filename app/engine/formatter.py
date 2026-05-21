"""排版引擎 — 读取 FormatConfig，对 .docx 文档应用格式"""

import copy
from typing import Optional, Callable, List

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.models.config import FormatConfig, ElementStyle

ALIGN_MAP = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class Formatter:
    def __init__(self, config: FormatConfig):
        self.cfg = config

    def format(self, input_path: str, output_path: str = None,
               progress: Optional[Callable[[str, int], None]] = None) -> str:
        """执行排版，返回输出路径"""
        if output_path is None:
            output_path = input_path

        def _progress(msg, pct):
            if progress:
                progress(msg, pct)

        _progress('正在加载文档...', 5)
        doc = Document(input_path)

        _progress('正在设置页面...', 10)
        self._setup_page(doc)

        _progress('正在分析文档结构...', 15)
        tags = self._classify_paragraphs(doc)
        tags = self._adjust_heading_levels(tags)

        _progress('正在应用段落格式...', 30)
        self._apply_paragraph_styles(doc, tags)

        _progress('正在格式化表格...', 55)
        self._format_tables(doc)

        _progress('正在设置分节...', 65)
        body_h1_indices, body_h1_texts = self._find_chapter_headings(doc)
        self._insert_section_breaks(doc, tags, body_h1_indices, body_h1_texts)

        _progress('正在插入目录...', 75)
        self._insert_toc(doc, tags)

        _progress('正在设置页眉页脚...', 85)
        self._setup_headers_footers(doc, tags, body_h1_indices, body_h1_texts)

        _progress('正在保存...', 95)
        doc.save(output_path)

        _progress('排版完成', 100)
        return output_path

    # ================================================================
    # 低层 OOXML 工具
    # ================================================================

    @staticmethod
    def _set_run_font(run, cn_font, en_font='Times New Roman', size=None,
                      bold=None, italic=None):
        if size is not None:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        run.font.name = en_font

        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)
        rFonts.set(qn('w:cs'), en_font)

    @staticmethod
    def _set_para_spacing(para, spacing, rule='multiple'):
        pf = para.paragraph_format
        if rule == 'multiple':
            pf.line_spacing = spacing
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        else:
            pf.line_spacing = Pt(spacing)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    @staticmethod
    def _set_first_line_indent(para, cm_val):
        para.paragraph_format.first_line_indent = Cm(cm_val)

    @staticmethod
    def _clear_first_line_indent(para):
        para.paragraph_format.first_line_indent = None

    @staticmethod
    def _clear_para_spacing(para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    def _apply_element_style(self, para, es: ElementStyle):
        """将 ElementStyle 应用到一个段落"""
        for run in para.runs:
            self._set_run_font(run, es.cn_font, es.en_font,
                              Pt(es.font_size), es.bold)
        self._set_para_spacing(para, es.line_spacing, es.line_spacing_rule)
        para.alignment = ALIGN_MAP.get(es.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        if es.first_line_indent is not None:
            self._set_first_line_indent(para, es.first_line_indent)
        else:
            self._clear_first_line_indent(para)
        self._clear_para_spacing(para)

    # ================================================================
    # 页面设置
    # ================================================================

    def _setup_page(self, doc):
        pg = self.cfg.page
        for section in doc.sections:
            section.page_width = Cm(pg.page_width)
            section.page_height = Cm(pg.page_height)
            section.top_margin = Cm(pg.top_margin)
            section.bottom_margin = Cm(pg.bottom_margin)
            section.left_margin = Cm(pg.left_margin)
            section.right_margin = Cm(pg.right_margin)

    # ================================================================
    # 段落分类
    # ================================================================

    def _classify_paragraphs(self, doc) -> List[str]:
        n = len(doc.paragraphs)
        tags = ['body'] * n
        toc_start = None
        ref_start = None
        appendix_start = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else 'Normal'

            if style == 'Heading 1':
                tags[i] = 'heading1'
            elif style == 'Heading 2':
                tags[i] = 'heading2'
            elif style == 'Heading 3':
                tags[i] = 'heading3'
            elif style == 'Heading 4':
                tags[i] = 'heading4'
            elif style == 'Source Code':
                tags[i] = 'source-code'

            if text == '摘要' or (text.startswith('摘要') and len(text) < 10):
                tags[i] = 'abstract-label'
            elif text == '关键词' or (text.startswith('关键词') and len(text) < 10):
                tags[i] = 'keywords-label'
            elif text.upper() == 'ABSTRACT' or (text.upper().startswith('ABSTRACT') and len(text) < 20):
                tags[i] = 'en-abstract-label'
            elif text.upper() in ('KEY WORDS', 'KEYWORDS') or (text.upper().startswith('KEY WORDS') and len(text) < 20):
                tags[i] = 'en-keywords-label'
            elif '目录' in text and len(text) <= 4:
                tags[i] = 'toc-heading'
                toc_start = i
            elif '参考文献' in text and len(text) <= 6:
                tags[i] = 'ref-heading'
                ref_start = i
            elif text == '附录' or (text.startswith('附录') and len(text) <= 6):
                tags[i] = 'appendix-heading'
                appendix_start = i
            elif text.startswith('附录 ') and len(text) > 6:
                tags[i] = 'appendix-sub'

        # 目录条目
        if toc_start is not None:
            for i in range(toc_start + 1, n):
                para = doc.paragraphs[i]
                text = para.text.strip()
                style = para.style.name if para.style else ''
                if style in ('Heading 1', 'Heading 2') and '目录' not in text:
                    break
                if text and style == 'Normal':
                    tags[i] = 'toc-item'

        # 参考文献条目
        if ref_start is not None:
            for i in range(ref_start + 1, n):
                para = doc.paragraphs[i]
                text = para.text.strip()
                style = para.style.name if para.style else ''
                if style in ('Heading 1', 'Heading 2') and '参考文献' not in text:
                    break
                if text and not text.startswith('附录'):
                    tags[i] = 'ref-item'

        # 附录子标题
        if appendix_start is not None:
            for i in range(appendix_start + 1, n):
                para = doc.paragraphs[i]
                style = para.style.name if para.style else ''
                if style == 'Heading 3' and '附录' in doc.paragraphs[i].text:
                    tags[i] = 'appendix-sub'

        # 中文摘要正文
        for i, t in enumerate(tags):
            if t == 'abstract-label':
                for j in range(i + 1, n):
                    para_j = doc.paragraphs[j]
                    text_j = para_j.text.strip()
                    style_j = para_j.style.name if para_j.style else ''
                    if text_j.startswith('关键词') or style_j in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'):
                        break
                    if tags[j] == 'body' and text_j:
                        tags[j] = 'abstract-body'
                break

        # 英文摘要正文
        for i, t in enumerate(tags):
            if t == 'en-abstract-label':
                for j in range(i + 1, n):
                    para_j = doc.paragraphs[j]
                    text_j = para_j.text.strip()
                    style_j = para_j.style.name if para_j.style else ''
                    if (text_j.upper().startswith('KEY WORDS') or
                        style_j in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4') or
                            tags[j] == 'abstract-label'):
                        break
                    if tags[j] == 'body' and text_j:
                        tags[j] = 'en-abstract-body'
                break

        # 图表题注
        body_ref_keywords = ['展示了', '可以看出', '从表', '从图', '给出了', '列出了', '显示了']
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if para.runs:
                is_caption = not any(kw in text for kw in body_ref_keywords)
                if is_caption:
                    if text.startswith('表 ') or text.startswith('表\n'):
                        tags[i] = 'table-caption'
                    elif text.startswith('图 ') or text.startswith('图\n'):
                        tags[i] = 'figure-caption'
                    elif text.startswith('表') and len(text) > 2:
                        tags[i] = 'table-caption'
                    elif text.startswith('图') and len(text) > 2:
                        tags[i] = 'figure-caption'

        # 文末注释
        note_keywords = ['需要我再提供', '注：文档部分内容', '（注：']
        for i in range(max(0, n - 5), n):
            text = doc.paragraphs[i].text.strip()
            if any(kw in text for kw in note_keywords):
                tags[i] = 'end-note'

        # 公式
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith("$") or text.startswith("x' =") or text.startswith("MAE"):
                if any(c in text for c in '∑\\frac'):
                    tags[i] = 'formula'

        # 题目页（摘要之前的内容）
        abstract_idx = -1
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ''
            if text == '摘要' and style.startswith('Heading'):
                abstract_idx = i
                break
        if abstract_idx < 0:
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text.startswith('摘要') and len(text) < 10:
                    abstract_idx = i
                    break

        if abstract_idx > 0:
            for i in range(abstract_idx):
                if tags[i] == 'heading1':
                    tags[i] = 'title-uni'
                elif tags[i] == 'heading2':
                    tags[i] = 'title-doctype'
                elif tags[i] == 'heading3':
                    tags[i] = 'title-paper'
                elif tags[i] == 'heading4':
                    tags[i] = 'title-date'
                elif tags[i] == 'body' and not doc.paragraphs[i].text.strip():
                    tags[i] = 'title-blank'
                elif tags[i] == 'body':
                    tags[i] = 'title-sep'

        return tags

    @staticmethod
    def _adjust_heading_levels(tags: List[str]) -> List[str]:
        has_heading1 = 'heading1' in tags
        if not has_heading1:
            for i in range(len(tags)):
                if tags[i] == 'heading2':
                    tags[i] = 'heading1'
                elif tags[i] == 'heading3':
                    tags[i] = 'heading2'
                elif tags[i] == 'heading4':
                    tags[i] = 'heading3'
        return tags

    # ================================================================
    # 段落格式应用
    # ================================================================

    def _apply_paragraph_styles(self, doc, tags):
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            tag = tags[i]

            # 题目页
            if tag == 'title-uni':
                self._apply_element_style(para, self.cfg.title_uni)
            elif tag == 'title-doctype':
                self._apply_element_style(para, self.cfg.title_doctype)
            elif tag == 'title-paper':
                self._apply_element_style(para, self.cfg.title_paper)
            elif tag == 'title-date':
                self._apply_element_style(para, self.cfg.title_date)
            elif tag == 'title-blank':
                self._clear_para_spacing(para)
                self._set_para_spacing(para, 12, 'exact')
            elif tag == 'title-sep':
                self._clear_para_spacing(para)

            # 一级标题
            elif tag == 'heading1' and '摘要' not in text:
                self._apply_element_style(para, self.cfg.heading_1)
                for run in para.runs:
                    self._set_run_font(run, self.cfg.heading_1.cn_font,
                                      self.cfg.heading_1.en_font,
                                      Pt(self.cfg.heading_1.font_size), False)
                # 一级标题前插入分页符
                pPr = para._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    para._element.insert(0, pPr)
                if pPr.find(qn('w:pageBreakBefore')) is None:
                    pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))

            # 摘要标签
            elif tag == 'abstract-label':
                self._apply_element_style(para, self.cfg.abstract_label)

            # 英文摘要标签
            elif tag == 'en-abstract-label':
                self._apply_element_style(para, self.cfg.en_abstract_label)

            # 目录标题
            elif tag == 'toc-heading':
                self._apply_element_style(para, self.cfg.toc.toc_heading_style)

            # 目录条目
            elif tag == 'toc-item':
                self._apply_element_style(para, self.cfg.toc.toc_item_style)

            # 二级标题
            elif tag == 'heading2':
                self._apply_element_style(para, self.cfg.heading_2)

            # 三级标题
            elif tag == 'heading3':
                self._apply_element_style(para, self.cfg.heading_3)

            # 四级标题
            elif tag == 'heading4':
                self._apply_element_style(para, self.cfg.heading_4)

            # 表题注
            elif tag == 'table-caption':
                for run in para.runs:
                    self._set_run_font(run, self.cfg.table_caption.cn_font,
                                      self.cfg.table_caption.en_font,
                                      Pt(self.cfg.table_caption.font_size), False)
                self._set_para_spacing(para, self.cfg.table_caption.line_spacing,
                                      self.cfg.table_caption.line_spacing_rule)
                para.alignment = ALIGN_MAP.get(self.cfg.table_caption.alignment,
                                              WD_ALIGN_PARAGRAPH.CENTER)
                self._clear_first_line_indent(para)
                self._clear_para_spacing(para)

            # 图题注
            elif tag == 'figure-caption':
                for run in para.runs:
                    self._set_run_font(run, self.cfg.figure_caption.cn_font,
                                      self.cfg.figure_caption.en_font,
                                      Pt(self.cfg.figure_caption.font_size), False)
                self._set_para_spacing(para, self.cfg.figure_caption.line_spacing,
                                      self.cfg.figure_caption.line_spacing_rule)
                para.alignment = ALIGN_MAP.get(self.cfg.figure_caption.alignment,
                                              WD_ALIGN_PARAGRAPH.CENTER)
                self._clear_first_line_indent(para)
                self._clear_para_spacing(para)

            # 参考文献标题
            elif tag == 'ref-heading':
                self._apply_element_style(para, self.cfg.ref_heading)

            # 参考文献条目
            elif tag == 'ref-item':
                self._apply_element_style(para, self.cfg.ref_item)

            # 附录大标题
            elif tag == 'appendix-heading':
                self._apply_element_style(para, self.cfg.appendix_heading)

            # 附录子标题
            elif tag == 'appendix-sub':
                self._apply_element_style(para, self.cfg.appendix_sub)

            # 源代码
            elif tag == 'source-code':
                self._apply_element_style(para, self.cfg.source_code)

            # 文末注释
            elif tag == 'end-note':
                self._apply_element_style(para, self.cfg.end_note)

            # 摘要/英文摘要正文
            elif tag in ('abstract-body', 'en-abstract-body'):
                es = self.cfg.abstract_body
                for run in para.runs:
                    if run.font.bold:
                        self._set_run_font(run, es.cn_font, es.en_font,
                                          Pt(es.font_size), True)
                    else:
                        self._set_run_font(run, es.cn_font, es.en_font,
                                          Pt(es.font_size), False)
                self._set_para_spacing(para, es.line_spacing, es.line_spacing_rule)
                para.alignment = ALIGN_MAP.get(es.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
                if para.paragraph_format.first_line_indent is None:
                    self._set_first_line_indent(para, es.first_line_indent or 0.74)
                self._clear_para_spacing(para)

            # 正文
            elif tag == 'body':
                es = self.cfg.body
                for run in para.runs:
                    if run.font.bold:
                        self._set_run_font(run, es.cn_font, es.en_font,
                                          Pt(es.font_size), True)
                    else:
                        self._set_run_font(run, es.cn_font, es.en_font,
                                          Pt(es.font_size), False)
                self._set_para_spacing(para, es.line_spacing, es.line_spacing_rule)
                para.alignment = ALIGN_MAP.get(es.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
                if para.paragraph_format.first_line_indent is None:
                    if es.first_line_indent is not None:
                        self._set_first_line_indent(para, es.first_line_indent)
                self._clear_para_spacing(para)

            # 公式
            elif tag == 'formula':
                es = self.cfg.formula
                for run in para.runs:
                    self._set_run_font(run, es.cn_font, es.en_font,
                                      Pt(es.font_size), False)
                self._set_para_spacing(para, es.line_spacing, es.line_spacing_rule)
                self._clear_para_spacing(para)
                self._clear_first_line_indent(para)
                tabs_xml = (
                    f'<w:tabs {nsdecls("w")}>'
                    f'<w:tab w:val="center" w:pos="4150"/>'
                    f'<w:tab w:val="right" w:pos="8301"/>'
                    f'</w:tabs>'
                )
                pPr = para._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    para._element.insert(0, pPr)
                existing = pPr.find(qn('w:tabs'))
                if existing is not None:
                    pPr.remove(existing)
                pPr.append(parse_xml(tabs_xml))
                if '\t' in para.text:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 特殊处理：摘要/关键词段落
        self._apply_special_labels(doc, tags)

    def _apply_special_labels(self, doc, tags):
        """处理摘要、关键词等特殊标签段落"""
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            tag = tags[i]

            if text.startswith('摘要') and tag in ('body', 'heading1'):
                self._apply_labeled_para(para, self.cfg.abstract_body, '摘要')
                tags[i] = 'abstract-body'

            if text.startswith('关键词') and tag == 'body':
                self._apply_labeled_para(para, self.cfg.keywords_body, '关键词')
                tags[i] = 'keywords-body'

            if text.upper().startswith('ABSTRACT') and tag == 'body':
                self._apply_labeled_para(para, self.cfg.en_abstract_body, 'ABSTRACT')
                tags[i] = 'en-abstract-body'

            if text.upper().startswith('KEY WORDS') and tag == 'body':
                self._apply_labeled_para(para, self.cfg.en_keywords_body, 'KEY WORDS')
                tags[i] = 'en-keywords-body'

    def _apply_labeled_para(self, para, es: ElementStyle, label: str):
        """格式化带标签的段落（摘要正文、关键词等）"""
        for run in para.runs:
            if label in run.text:
                self._set_run_font(run, es.cn_font, es.en_font,
                                  Pt(es.font_size), True)
            else:
                self._set_run_font(run, es.cn_font, es.en_font,
                                  Pt(es.font_size), False)
        self._set_para_spacing(para, es.line_spacing, es.line_spacing_rule)
        para.alignment = ALIGN_MAP.get(es.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        if para.paragraph_format.first_line_indent is None and es.first_line_indent is not None:
            self._set_first_line_indent(para, es.first_line_indent)
        self._clear_para_spacing(para)

    # ================================================================
    # 表格格式化
    # ================================================================

    def _format_tables(self, doc):
        tsc = self.cfg.table_style
        title_page_fields = {'学院', '专业', '姓名', '学号', '指导教师',
                             '学　　院', '专　　业', '姓　　名', '学　　号'}

        for idx, table in enumerate(doc.tables):
            is_title_table = False
            if idx == 0 and len(table.rows) > 0:
                first_cell_text = ''
                if table.rows[0].cells:
                    first_cell_text = table.rows[0].cells[0].text.strip()
                is_title_table = (first_cell_text in title_page_fields or
                                  '学院' in first_cell_text)

            if is_title_table:
                ett = self.cfg.title_table
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.font.bold or '**' in run.text:
                                    self._set_run_font(run, ett.cn_font, ett.en_font,
                                                      Pt(ett.font_size), True)
                                else:
                                    self._set_run_font(run, ett.cn_font, ett.en_font,
                                                      Pt(ett.font_size), False)
                            self._set_para_spacing(para, ett.line_spacing,
                                                  ett.line_spacing_rule)
                            para.alignment = ALIGN_MAP.get(ett.alignment,
                                                          WD_ALIGN_PARAGRAPH.CENTER)
                            self._clear_first_line_indent(para)
                            self._clear_para_spacing(para)
                self._remove_table_borders(table)
            else:
                # 三线表
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                self._set_run_font(run, tsc.cn_font, tsc.en_font,
                                                  Pt(tsc.font_size), False)
                            self._set_para_spacing(para, tsc.line_spacing, 'multiple')
                            para.alignment = ALIGN_MAP.get(tsc.alignment,
                                                          WD_ALIGN_PARAGRAPH.CENTER)
                            self._clear_first_line_indent(para)
                            self._clear_para_spacing(para)
                self._remove_table_borders(table)
                top_sz = str(int(tsc.top_border_width * 8))
                bottom_sz = str(int(tsc.bottom_border_width * 8))
                header_sz = str(int(tsc.header_bottom_width * 8))
                self._set_table_border(table, 'top', top_sz)
                self._set_table_border(table, 'bottom', bottom_sz)
                if len(table.rows) > 0:
                    for cell in table.rows[0].cells:
                        self._set_cell_bottom_border(cell, header_sz)

    @staticmethod
    def _remove_table_borders(table):
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            table._tbl.insert(0, tblPr)
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)

    @staticmethod
    def _set_table_border(table, position, size, color='000000'):
        tblPr = table._tbl.find(qn('w:tblPr'))
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is None:
            borders = parse_xml(f'<w:tblBorders {nsdecls("w")}></w:tblBorders>')
            tblPr.append(borders)
        border_el = parse_xml(
            f'<w:{position} {nsdecls("w")} '
            f'w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        )
        borders.append(border_el)

    @staticmethod
    def _set_cell_bottom_border(cell, size, color='000000'):
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            cell._tc.insert(0, tcPr)
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        existing = tcPr.find(qn('w:tcBorders'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcBorders)

    # ================================================================
    # 分节与页眉页脚
    # ================================================================

    def _find_chapter_headings(self, doc):
        """查找正文章节标题"""
        content_start = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text == '摘要' or text == 'ABSTRACT' or '目录' in text:
                content_start = i + 1
                break

        body_h1_indices = []
        body_h1_texts = []
        for i, para in enumerate(doc.paragraphs):
            if i < content_start:
                continue
            text = para.text.strip()
            style = para.style.name if para.style else ''
            if (style == 'Heading 1' and '摘要' not in text and
                    'ABSTRACT' not in text and '目录' not in text):
                body_h1_indices.append(i)
                body_h1_texts.append(text)

        return body_h1_indices, body_h1_texts

    def _insert_section_breaks(self, doc, tags, body_h1_indices, body_h1_texts):
        # 检测题目页
        title_page_indices = [i for i, t in enumerate(tags) if t.startswith('title-')]
        has_title_page = len(title_page_indices) > 0

        # 检测摘要
        cn_ab_tag = [i for i, t in enumerate(tags) if t in ('abstract-body', 'abstract-label')]
        cn_kw_tag = [i for i, t in enumerate(tags) if t in ('keywords-body', 'keywords-label')]
        en_ab_tag = [i for i, t in enumerate(tags) if t in ('en-abstract-body', 'en-abstract-label')]
        en_kw_tag = [i for i, t in enumerate(tags) if t in ('en-keywords-body', 'en-keywords-label')]
        has_abstract = len(cn_ab_tag) > 0
        has_en_abstract = len(en_ab_tag) > 0

        if has_title_page:
            last_title = max(title_page_indices)
            self._insert_section_break(doc.paragraphs[last_title])

        if has_abstract:
            last_cn = max(cn_ab_tag + cn_kw_tag)
            self._insert_section_break(doc.paragraphs[last_cn])

        if has_en_abstract:
            last_en = max(en_ab_tag + en_kw_tag)
            self._insert_section_break(doc.paragraphs[last_en])

        if len(body_h1_indices) >= 1:
            for h1_idx in reversed(body_h1_indices[:]):
                prev_para = doc.paragraphs[h1_idx - 1]
                self._insert_section_break(prev_para)

    def _insert_section_break(self, para):
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para._element.insert(0, pPr)
        if pPr.find(qn('w:sectPr')) is not None:
            return
        body = para._element.getparent()
        body_sectPr = body.find(qn('w:sectPr'))
        if body_sectPr is not None:
            pPr.append(copy.deepcopy(body_sectPr))
            new_sectPr = pPr.find(qn('w:sectPr'))
            if new_sectPr is not None:
                for tn in ('w:headerReference', 'w:footerReference',
                           'w:evenAndOddHeaders', 'w:titlePg'):
                    for el in new_sectPr.findall(qn(tn)):
                        new_sectPr.remove(el)
            for tn in ('w:headerReference', 'w:footerReference',
                       'w:evenAndOddHeaders', 'w:titlePg'):
                for el in body_sectPr.findall(qn(tn)):
                    body_sectPr.remove(el)

    def _insert_toc(self, doc, tags):
        if not self.cfg.toc.auto_insert:
            return

        cn_kw_tag = [i for i, t in enumerate(tags) if t in ('keywords-body', 'keywords-label')]
        en_kw_tag = [i for i, t in enumerate(tags) if t in ('en-keywords-body', 'en-keywords-label')]
        cn_ab_tag = [i for i, t in enumerate(tags) if t in ('abstract-body', 'abstract-label')]
        en_ab_tag = [i for i, t in enumerate(tags) if t in ('en-abstract-body', 'en-abstract-label')]

        all_kw = en_kw_tag + cn_kw_tag
        all_ab = en_ab_tag + cn_ab_tag

        toc_insert_idx = None
        if all_kw:
            toc_insert_idx = max(all_kw) + 1
        elif all_ab:
            toc_insert_idx = max(all_ab) + 1

        if toc_insert_idx is None:
            for i, t in enumerate(tags):
                if t == 'heading1' and '摘要' not in doc.paragraphs[i].text:
                    toc_insert_idx = i
                    break

        if toc_insert_idx is None:
            return

        # 防重复
        for p in doc.paragraphs:
            for instr in p._element.findall('.//' + qn('w:instrText')):
                if instr.text and 'TOC' in instr.text:
                    return

        tocs = self.cfg.toc
        ref_element = doc.paragraphs[toc_insert_idx]._element
        body = doc.element.body

        # 目录标题
        th = tocs.toc_heading_style
        line_val = str(int(th.line_spacing * 240))
        toc_heading_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        toc_heading_pPr = parse_xml(
            f'<w:pPr {nsdecls("w")}>'
            f'<w:jc w:val="center"/>'
            f'<w:spacing w:line="{line_val}" w:lineRule="auto" w:before="0" w:after="0"/>'
            f'</w:pPr>'
        )
        toc_heading_p.insert(0, toc_heading_pPr)
        toc_heading_r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        toc_heading_rPr = parse_xml(
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:eastAsia="{th.cn_font}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'<w:sz w:val="{int(th.font_size * 2)}"/>'
            f'</w:rPr>'
        )
        toc_heading_r.insert(0, toc_heading_rPr)
        toc_heading_t = parse_xml(
            f'<w:t {nsdecls("w")} xml:space="preserve">目录</w:t>'
        )
        toc_heading_r.append(toc_heading_t)
        toc_heading_p.append(toc_heading_r)

        # TOC 域
        ti = tocs.toc_item_style
        ti_line = str(int(ti.line_spacing * 240))
        toc_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        toc_pPr = parse_xml(
            f'<w:pPr {nsdecls("w")}>'
            f'<w:spacing w:line="{ti_line}" w:lineRule="auto" w:before="0" w:after="0"/>'
            f'</w:pPr>'
        )
        toc_p.insert(0, toc_pPr)

        r_begin = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r_begin.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        toc_p.append(r_begin)

        r_instr = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        levels = tocs.heading_levels
        r_instr.append(
            parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "{levels}" \\h \\z \\u </w:instrText>')
        )
        toc_p.append(r_instr)

        r_sep = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r_sep.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
        toc_p.append(r_sep)

        r_hint = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r_hint_rPr = parse_xml(
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:eastAsia="宋体" w:ascii="Times New Roman"/>'
            f'<w:sz w:val="21"/>'
            f'</w:rPr>'
        )
        r_hint.insert(0, r_hint_rPr)
        r_hint.append(
            parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">（请在 Word 中右键此处 → 更新域，以生成目录）</w:t>')
        )
        toc_p.append(r_hint)

        r_end = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        r_end.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        toc_p.append(r_end)

        ref_element.addprevious(toc_p)
        ref_element.addprevious(toc_heading_p)

    # ================================================================
    # 页眉页脚设置
    # ================================================================

    def _setup_headers_footers(self, doc, tags, body_h1_indices, body_h1_texts):
        hf = self.cfg.header_footer

        # 检测各区域
        title_page_indices = [i for i, t in enumerate(tags) if t.startswith('title-')]
        has_title_page = len(title_page_indices) > 0
        cn_ab_tag = [i for i, t in enumerate(tags) if t in ('abstract-body', 'abstract-label')]
        cn_kw_tag = [i for i, t in enumerate(tags) if t in ('keywords-body', 'keywords-label')]
        en_ab_tag = [i for i, t in enumerate(tags) if t in ('en-abstract-body', 'en-abstract-label')]
        en_kw_tag = [i for i, t in enumerate(tags) if t in ('en-keywords-body', 'en-keywords-label')]
        has_abstract = len(cn_ab_tag) > 0
        has_en_abstract = len(en_ab_tag) > 0

        # 确定奇数页页眉文字（用于摘要节）
        def _abstract_header(idx, title_offset, abs_offset):
            if idx < abs_offset - 1:
                if has_abstract and has_en_abstract:
                    return '摘要' if idx - title_offset == 0 else 'ABSTRACT'
                elif has_abstract:
                    return '摘要'
                elif has_en_abstract:
                    return 'ABSTRACT'
                else:
                    return '摘要'
            else:
                return '目录'

        title_offset = 1 if has_title_page else 0
        abs_offset = title_offset
        if has_abstract:
            abs_offset += 1
        if has_en_abstract:
            abs_offset += 1
        abs_offset += 1  # TOC 节

        roman_started = False
        decimal_started = False

        for idx, section in enumerate(doc.sections):
            if has_title_page and idx == 0:
                continue

            is_pre_content = idx < abs_offset

            if is_pre_content:
                hdr = _abstract_header(idx, title_offset, abs_offset)
                if not roman_started:
                    self._setup_section_header_footer(
                        section, hdr, hdr, 'upperRoman', 1)
                    roman_started = True
                else:
                    self._setup_section_header_footer(
                        section, hdr, hdr, 'upperRoman')
            else:
                chapter_idx = idx - abs_offset
                if chapter_idx < len(body_h1_texts):
                    chap_title = body_h1_texts[chapter_idx]
                else:
                    chap_title = hf.even_header_text
                if hf.odd_header_source == 'fixed' and hf.odd_header_fixed_text:
                    chap_title = hf.odd_header_fixed_text
                if not decimal_started:
                    self._setup_section_header_footer(
                        section, chap_title, hf.even_header_text,
                        hf.page_num_format, 1)
                    decimal_started = True
                else:
                    self._setup_section_header_footer(
                        section, chap_title, hf.even_header_text,
                        hf.page_num_format)

    def _setup_section_header_footer(self, section, odd_text, even_text,
                                     page_num_format='decimal', page_start=None):
        hf = self.cfg.header_footer

        sectPr = section._sectPr

        # 奇偶页不同
        if sectPr.find(qn('w:evenAndOddHeaders')) is None:
            sectPr.append(parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>'))

        # 页码格式
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}/>')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:fmt'), page_num_format)
        if page_start is not None:
            pgNumType.set(qn('w:start'), str(page_start))

        font_size = Pt(hf.font_size)

        # 偶数页页眉
        even_hdr = section.even_page_header
        even_hdr.is_linked_to_previous = False
        for p in even_hdr.paragraphs:
            p.clear()
        hp_even = even_hdr.paragraphs[0] if even_hdr.paragraphs else even_hdr.add_paragraph()
        hp_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp_even.add_run(even_text)
        self._set_run_font(run, hf.cn_font, hf.en_font, font_size)
        if hf.header_line:
            self._add_header_line(hp_even)

        # 偶数页页脚
        even_ftr = section.even_page_footer
        even_ftr.is_linked_to_previous = False
        for p in even_ftr.paragraphs:
            p.clear()
        fp_even = even_ftr.paragraphs[0] if even_ftr.paragraphs else even_ftr.add_paragraph()
        fp_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_page_number(fp_even, hf.cn_font, font_size, hf.en_font)

        # 奇数页页眉
        odd_hdr = section.header
        odd_hdr.is_linked_to_previous = False
        for p in odd_hdr.paragraphs:
            p.clear()
        hp_odd = odd_hdr.paragraphs[0] if odd_hdr.paragraphs else odd_hdr.add_paragraph()
        hp_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp_odd.add_run(odd_text)
        self._set_run_font(run, hf.cn_font, hf.en_font, font_size)
        if hf.header_line:
            self._add_header_line(hp_odd)

        # 奇数页页脚
        odd_ftr = section.footer
        odd_ftr.is_linked_to_previous = False
        for p in odd_ftr.paragraphs:
            p.clear()
        fp_odd = odd_ftr.paragraphs[0] if odd_ftr.paragraphs else odd_ftr.add_paragraph()
        fp_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_page_number(fp_odd, hf.cn_font, font_size, hf.en_font)

    def _add_page_number(self, paragraph, cn_font, font_size, en_font='Times New Roman'):
        run_b = paragraph.add_run()
        run_b._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run_i = paragraph.add_run()
        run_i._element.append(
            parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run_e = paragraph.add_run()
        run_e._element.append(
            parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        for r in paragraph.runs:
            self._set_run_font(r, cn_font, en_font, font_size)

    def _add_header_line(self, paragraph):
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            paragraph._element.insert(0, pPr)
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
            f'</w:pBdr>'
        )
        existing = pPr.find(qn('w:pBdr'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(pBdr)
