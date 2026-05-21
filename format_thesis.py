"""
河南理工大学本科毕业设计（论文）排版脚本
============================================
根据《河南理工大学本科毕业设计（论文）撰写规范》对 test.docx 进行自动排版。

使用内容检测代替固定段落索引，适配 pandoc 从 markdown 生成的文档。

中文字号对照：
  三号=16pt, 小三=15pt, 四号=14pt, 小四=12pt, 五号=10.5pt, 小五=9pt

依赖：python-docx
"""

import sys
import re
import copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = "《地球物理综合应用与实践》报告.docx"
doc = Document(SRC)


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, cn_font, en_font='Times New Roman', size=None, bold=None, italic=None):
    """设置文字段（run）的中英文字体、大小、加粗/斜体"""
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic

    run.font.name = en_font

    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._element.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)


def set_para_spacing(para, spacing):
    """设置段落行距（float=多倍行距, Pt=固定值）"""
    pf = para.paragraph_format
    if isinstance(spacing, float):
        pf.line_spacing = spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    else:
        pf.line_spacing = spacing
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def clear_first_line_indent(para):
    """清除首行缩进"""
    para.paragraph_format.first_line_indent = None


def set_first_line_indent(para, cm=0.74):
    """设置首行缩进（默认 0.74cm ≈ 2 个中文字符）"""
    para.paragraph_format.first_line_indent = Cm(cm)


def clear_para_spacing(para):
    """清除段前段后间距"""
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def format_runs_in_para(para, cn_font, en_font='Times New Roman', size=None, bold=None):
    """对段落中所有 run 统一设置字体属性"""
    for run in para.runs:
        set_run_font(run, cn_font, en_font, size, bold)


def format_para_full(para, cn_font, size, bold, alignment, line_spacing,
                     first_line_indent=None, en_font='Times New Roman'):
    """一次性设置段落所有格式"""
    format_runs_in_para(para, cn_font, en_font, size, bold)
    set_para_spacing(para, line_spacing)
    para.alignment = alignment
    if first_line_indent is not None:
        set_first_line_indent(para, first_line_indent)
    else:
        clear_first_line_indent(para)
    clear_para_spacing(para)


# ============================================================
# 第一步：页面设置 → A4 纸 + 规范边距
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


# ============================================================
# 第二步：段落分类 → 识别每个段落的角色
# ============================================================
# 通过内容和样式判断段落类型，避免依赖硬编码索引

def classify_paragraphs(doc):
    """
    分析文档结构，返回每个段落的状态标记：
      'title'       — 论文标题行（正文最开头的一级标题）
      'abstract'    — 摘要
      'keywords'    — 关键词
      'toc-heading' — 目录标题
      'toc-item'    — 目录条目
      'ref-heading' — 参考文献标题
      'ref-item'    — 参考文献条目
      'appendix-heading' — 附录大标题
      'appendix-sub' — 附录子标题
      'source-code' — 源代码
      'heading1'    — 一级标题
      'heading2'    — 二级标题
      'heading3'    — 三级标题
      'heading4'    — 四级标题
      'caption'     — 图表题注
      'body'        — 正文段落
      'formula'     — 公式行
      'end-note'    — 文末注释
    """
    n = len(doc.paragraphs)
    tags = ['body'] * n

    # 先收集所有标题的位置和级别
    toc_start = None
    toc_end = None
    ref_start = None
    ref_end = None
    appendix_start = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else 'Normal'

        # --- 识别标题样式 ---
        if style == 'Heading 1':
            tags[i] = 'heading1'
        elif style == 'Heading 2':
            tags[i] = 'heading2'
        elif style == 'Heading 3':
            tags[i] = 'heading3'
        elif style == 'Heading 4':
            tags[i] = 'heading4'
        elif style == 'Source Code':
            tags[i] = 'source-code'

        # --- 识别特殊内容 ---
        if text == '摘要' or (text.startswith('摘要') and len(text) < 10):
            tags[i] = 'abstract-label'
        elif text == '关键词' or (text.startswith('关键词') and len(text) < 10):
            tags[i] = 'keywords-label'
        elif text.upper() == 'ABSTRACT' or (text.upper().startswith('ABSTRACT') and len(text) < 20):
            tags[i] = 'en-abstract-label'
        elif text.upper() in ('KEY WORDS', 'KEYWORDS') or (text.upper().startswith('KEY WORDS') and len(text) < 20):
            tags[i] = 'en-keywords-label'
        elif '目录' in text and len(text) <= 4:
            tags[i] = 'toc-heading'
            toc_start = i
        elif '参考文献' in text and len(text) <= 6:
            tags[i] = 'ref-heading'
            ref_start = i
        elif text == '附录' or (text.startswith('附录') and len(text) <= 6):
            tags[i] = 'appendix-heading'
            appendix_start = i
        elif text.startswith('附录 ') and len(text) > 6:
            tags[i] = 'appendix-sub'

    # --- 确定目录条目的范围 ---
    if toc_start is not None:
        # 目录条目：目录标题之后，直到下一个 Heading 1/2 章节标题
        for i in range(toc_start + 1, n):
            para = doc.paragraphs[i]
            text = para.text.strip()
            style = para.style.name if para.style else ''
            # 遇到下一个大标题就停止
            if style in ('Heading 1', 'Heading 2') and '目录' not in text:
                break
            if text and style == 'Normal':
                tags[i] = 'toc-item'
            toc_end = i

    # --- 确定参考文献条目的范围 ---
    if ref_start is not None:
        for i in range(ref_start + 1, n):
            para = doc.paragraphs[i]
            text = para.text.strip()
            style = para.style.name if para.style else ''
            # 遇到附录或下一章节就停止
            if style in ('Heading 1', 'Heading 2') and '参考文献' not in text:
                break
            if text and not text.startswith('附录'):
                tags[i] = 'ref-item'
            ref_end = i

    # --- 确定附录条目的范围 ---
    if appendix_start is not None:
        for i in range(appendix_start + 1, n):
            para = doc.paragraphs[i]
            style = para.style.name if para.style else ''
            if style == 'Heading 3' and '附录' in doc.paragraphs[i].text:
                tags[i] = 'appendix-sub'

    # --- 确定中文摘要正文范围（摘要标题之后、关键词之前的所有正文段落） ---
    for i, t in enumerate(tags):
        if t == 'abstract-label':
            for j in range(i + 1, n):
                para_j = doc.paragraphs[j]
                text_j = para_j.text.strip()
                style_j = para_j.style.name if para_j.style else ''
                # 遇到关键词文本、任何标题则停止
                if text_j.startswith('关键词') or style_j in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'):
                    break
                if tags[j] == 'body' and text_j:
                    tags[j] = 'abstract-body'
            break

    # --- 确定英文摘要正文范围（ABSTRACT 标题之后、KEY WORDS 之前的所有正文段落） ---
    for i, t in enumerate(tags):
        if t == 'en-abstract-label':
            for j in range(i + 1, n):
                para_j = doc.paragraphs[j]
                text_j = para_j.text.strip()
                style_j = para_j.style.name if para_j.style else ''
                if text_j.upper().startswith('KEY WORDS') or style_j in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4') or tags[j] == 'abstract-label':
                    break
                if tags[j] == 'body' and text_j:
                    tags[j] = 'en-abstract-body'
            break

    # --- 识别正文中的图表题注 ---
    body_ref_keywords = ['展示了', '可以看出', '从表', '从图', '给出了', '列出了', '显示了']
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if para.runs:
            is_caption = not any(kw in text for kw in body_ref_keywords)
            if is_caption:
                if text.startswith('表 ') or text.startswith('表\n'):
                    tags[i] = 'table-caption'
                elif text.startswith('图 ') or text.startswith('图\n'):
                    tags[i] = 'figure-caption'
                elif text.startswith('表') and len(text) > 2:
                    tags[i] = 'table-caption'
                elif text.startswith('图') and len(text) > 2:
                    tags[i] = 'figure-caption'

    # --- 识别文末注释（文档末尾的"需要我再提供..."和"注：..."等） ---
    note_keywords = ['需要我再提供', '注：文档部分内容', '（注：']
    for i in range(max(0, n - 5), n):
        text = doc.paragraphs[i].text.strip()
        if any(kw in text for kw in note_keywords):
            tags[i] = 'end-note'

    # --- 识别公式行（以 $ 开头或包含 LaTeX 公式语法） ---
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("$") or text.startswith("x' =") or text.startswith("MAE"):
            if any(c in text for c in '∑\\frac'):
                tags[i] = 'formula'

    # --- 检测题目页（"摘要"之前的所有内容） ---
    abstract_idx = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ''
        if text == '摘要' and style.startswith('Heading'):
            abstract_idx = i
            break
    # fallback: 查找第一个内容章节标题
    if abstract_idx < 0:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith('摘要') and len(text) < 10:
                abstract_idx = i
                break

    if abstract_idx > 0:
        for i in range(abstract_idx):
            if tags[i] == 'heading1':
                tags[i] = 'title-uni'
            elif tags[i] == 'heading2':
                tags[i] = 'title-doctype'
            elif tags[i] == 'heading3':
                tags[i] = 'title-paper'
            elif tags[i] == 'heading4':
                tags[i] = 'title-date'
            elif tags[i] == 'body' and not doc.paragraphs[i].text.strip():
                tags[i] = 'title-blank'
            elif tags[i] == 'body':
                # 题目页与摘要之间的分隔线、空白等
                tags[i] = 'title-sep'

    return tags


tags = classify_paragraphs(doc)

# ============================================================
# 自适应标题层级：如无 Heading 1，则 Heading 2/3 各升一级
# ============================================================
has_heading1 = 'heading1' in tags
if not has_heading1:
    # H2 → 一级, H3 → 二级, H4 → 三级
    for i in range(len(tags)):
        if tags[i] == 'heading2':
            tags[i] = 'heading1'
        elif tags[i] == 'heading3':
            tags[i] = 'heading2'
        elif tags[i] == 'heading4':
            tags[i] = 'heading3'

# ============================================================
# 第三步：逐段应用格式
# ============================================================

# 论文标题 — 收集所有标题页的连续 heading1 段落
title_indices = [i for i, t in enumerate(tags) if t == 'heading1']
# "摘要"可能被标记为 heading1 或 abstract-label
abstract_label_idx = None
keywords_label_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    tag = tags[i]
    style = para.style.name if para.style else 'Normal'

    # ---- 题目页 - 大学名称：小初号(36pt)黑体，居中 ----
    if tag == 'title-uni':
        format_para_full(para, '黑体', Pt(36), False, WD_ALIGN_PARAGRAPH.CENTER, 1.5)

    # ---- 题目页 - 文档类型：二号(22pt)黑体，居中 ----
    elif tag == 'title-doctype':
        format_para_full(para, '黑体', Pt(22), False, WD_ALIGN_PARAGRAPH.CENTER, 1.5)

    # ---- 题目页 - 论文题目：小二号(18pt)黑体，居中 ----
    elif tag == 'title-paper':
        format_para_full(para, '黑体', Pt(18), False, WD_ALIGN_PARAGRAPH.CENTER, 1.5)

    # ---- 题目页 - 日期：小三(15pt)宋体，居中 ----
    elif tag == 'title-date':
        format_para_full(para, '宋体', Pt(15), False, WD_ALIGN_PARAGRAPH.CENTER, 1.5)

    # ---- 题目页 - 空白行 ----
    elif tag == 'title-blank':
        clear_para_spacing(para)
        set_para_spacing(para, Pt(12))

    # ---- 题目页 - 分隔线 ----
    elif tag == 'title-sep':
        clear_para_spacing(para)

    # ---- 一级标题 (Heading 1)：小三黑体，居中，2 倍行距，每章另起一页 ----
    elif tag == 'heading1' and '摘要' not in text:
        format_para_full(para, '黑体', Pt(15), False, WD_ALIGN_PARAGRAPH.CENTER, 2.0)
        for run in para.runs:
            set_run_font(run, '黑体', 'Times New Roman', Pt(15), False)
        # 在一级标题前插入分页符（章另起一页）
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para._element.insert(0, pPr)
        pageBreak = pPr.find(qn('w:pageBreakBefore'))
        if pageBreak is None:
            pageBreak = parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>')
            pPr.append(pageBreak)

    # ---- 摘要标签 ("摘要")：小三黑体，居中，2 倍行距 ----
    elif tag == 'abstract-label':
        format_para_full(para, '黑体', Pt(15), False, WD_ALIGN_PARAGRAPH.CENTER, 2.0)
        abstract_label_idx = i

    # ---- 关键词标签 ("关键词") ----
    elif tag == 'keywords-label':
        keywords_label_idx = i

    # ---- 英文摘要标签 ("ABSTRACT")：小三，Times New Roman 加粗或黑体居中，2 倍行距 ----
    elif tag == 'en-abstract-label':
        format_para_full(para, '黑体', Pt(15), True, WD_ALIGN_PARAGRAPH.CENTER, 2.0)

    # ---- 英文关键词标签 ("KEY WORDS") ----
    elif tag == 'en-keywords-label':
        pass  # 与英文摘要同页，格式在后续特殊处理中应用

    # ---- 目录标题：小三黑体，居中，2 倍行距 ----
    elif tag == 'toc-heading':
        format_para_full(para, '黑体', Pt(15), False, WD_ALIGN_PARAGRAPH.CENTER, 2.0)

    # ---- 目录条目：小四宋体，左对齐，1.3 倍行距，无缩进 ----
    elif tag == 'toc-item':
        format_para_full(para, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 1.3)

    # ---- 二级标题：四号宋体加粗，左对齐，1.5 倍行距 ----
    elif tag == 'heading2':
        format_para_full(para, '宋体', Pt(14), True, WD_ALIGN_PARAGRAPH.LEFT, 1.5)

    # ---- 三级标题：小四宋体，首行缩进 2 字符，1.3 倍行距 ----
    elif tag == 'heading3':
        format_para_full(para, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 1.3,
                         first_line_indent=0.74)

    # ---- 四级标题：同三级 ----
    elif tag == 'heading4':
        format_para_full(para, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 1.3,
                         first_line_indent=0.74)

    # ---- 表题注：五号宋体，单倍行距，居中（规范：表名在表格上方正中，序号在表名左方） ----
    elif tag == 'table-caption':
        for run in para.runs:
            set_run_font(run, '宋体', 'Times New Roman', Pt(10.5), False)
        set_para_spacing(para, 1.0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_first_line_indent(para)
        clear_para_spacing(para)

    # ---- 图题注：五号宋体，单倍行距，居中（规范：图名在下方居中，序号在图名左方） ----
    elif tag == 'figure-caption':
        for run in para.runs:
            set_run_font(run, '宋体', 'Times New Roman', Pt(10.5), False)
        set_para_spacing(para, 1.0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_first_line_indent(para)
        clear_para_spacing(para)

    # ---- 参考文献标题：小三黑体，居中，2 倍行距 ----
    elif tag == 'ref-heading':
        format_para_full(para, '黑体', Pt(15), False, WD_ALIGN_PARAGRAPH.CENTER, 2.0)

    # ---- 参考文献条目：五号宋体，两端对齐，单倍行距 ----
    elif tag == 'ref-item':
        format_para_full(para, '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.0)

    # ---- 附录大标题：小三黑体，居中，2 倍行距 ----
    elif tag == 'appendix-heading':
        format_para_full(para, '黑体', Pt(15), False, WD_ALIGN_PARAGRAPH.CENTER, 2.0)

    # ---- 附录子标题：按二/三级标题处理 ----
    elif tag == 'appendix-sub':
        format_para_full(para, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 1.3,
                         first_line_indent=0.74)

    # ---- 源代码：小五宋体（9pt），左对齐，单倍行距 ----
    elif tag == 'source-code':
        format_para_full(para, '宋体', Pt(9), False, WD_ALIGN_PARAGRAPH.LEFT, 1.0)

    # ---- 文末注释：五号宋体，单倍行距 ----
    elif tag == 'end-note':
        format_para_full(para, '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.LEFT, 1.0)

    # ---- 摘要/英文摘要正文：同正文格式（小四宋体，两端对齐，首行缩进 2 字符，1.3 倍行距） ----
    elif tag in ('abstract-body', 'en-abstract-body'):
        for run in para.runs:
            if run.font.bold:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)
            else:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.paragraph_format.first_line_indent is None:
            set_first_line_indent(para, 0.74)
        clear_para_spacing(para)

    # ---- 正文：小四宋体，两端对齐，首行缩进 2 字符，1.3 倍行距 ----
    elif tag == 'body':
        # 保留特殊加粗（如摘要正文段落的"摘要"关键词等）
        for run in para.runs:
            if run.font.bold:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)
            else:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.paragraph_format.first_line_indent is None:
            set_first_line_indent(para, 0.74)
        clear_para_spacing(para)

    # ---- 公式：居中，编号右对齐（制表位：中 7.32cm + 右 14.64cm） ----
    elif tag == 'formula':
        format_runs_in_para(para, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        clear_para_spacing(para)
        clear_first_line_indent(para)
        # 添加制表位：居中位（页面中心）和右对齐位（右边界）
        # 可用宽度 = 21 - 3.18 - 3.18 = 14.64cm，1cm = 567 twips
        tabs_xml = (
            f'<w:tabs {nsdecls("w")}>'
            f'<w:tab w:val="center" w:pos="4150"/>'
            f'<w:tab w:val="right" w:pos="8301"/>'
            f'</w:tabs>'
        )
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para._element.insert(0, pPr)
        existing = pPr.find(qn('w:tabs'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(parse_xml(tabs_xml))
        # 含制表符的公式（tab+公式+tab+编号）用左对齐驱动制表位，否则直接居中
        if '\t' in para.text:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---- 特殊处理：摘要页面的正文段落 ----
# 摘要下的第一段正文通常以"摘要"开头，需要加粗"摘要"标签
# 关键词段落需要加粗"关键词"标签
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    tag = tags[i]

    # 摘要正文：查找"摘要"开头的段落（紧跟摘要标签之后）
    if text.startswith('摘要') and tag in ('body', 'heading1'):
        # 这是摘要正文 — 小四宋体，加粗"摘要"，其余正常
        for run in para.runs:
            if '摘要' in run.text:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)
            else:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.paragraph_format.first_line_indent is None:
            set_first_line_indent(para, 0.74)
        clear_para_spacing(para)
        tags[i] = 'abstract-body'

    # 关键词正文：查找"关键词"开头的段落
    if text.startswith('关键词') and tag in ('body',):
        for run in para.runs:
            if '关键词' in run.text:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)
            else:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.paragraph_format.first_line_indent is None:
            set_first_line_indent(para, 0.74)
        clear_para_spacing(para)
        tags[i] = 'keywords-body'

    # 英文摘要正文：查找"ABSTRACT"开头的段落
    if text.upper().startswith('ABSTRACT') and tag in ('body',):
        for run in para.runs:
            if 'ABSTRACT' in run.text.upper():
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)
            else:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.paragraph_format.first_line_indent is None:
            set_first_line_indent(para, 0.74)
        clear_para_spacing(para)
        tags[i] = 'en-abstract-body'

    # 英文关键词正文：查找"KEY WORDS"开头的段落
    if text.upper().startswith('KEY WORDS') and tag in ('body',):
        for run in para.runs:
            if 'KEY WORDS' in run.text.upper():
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)
            else:
                set_run_font(run, '宋体', 'Times New Roman', Pt(12), False)
        set_para_spacing(para, 1.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.paragraph_format.first_line_indent is None:
            set_first_line_indent(para, 0.74)
        clear_para_spacing(para)
        tags[i] = 'en-keywords-body'


# ============================================================
# 第四步：确定页眉文字（在插入目录前获取，避免索引偏移）
# ============================================================

odd_header_text = '河南理工大学本科毕业设计（论文）'
for i, para in enumerate(doc.paragraphs):
    t = tags[i]
    text = para.text.strip()
    if t == 'heading1' and '摘要' not in text:
        odd_header_text = text
        break
else:
    for i, para in enumerate(doc.paragraphs):
        t = tags[i]
        text = para.text.strip()
        if t == 'heading2':
            odd_header_text = text
            break

even_header_text = '河南理工大学本科毕业设计（论文）'

# 检测中/英文摘要（供 TOC 插入和后续分节使用）
cn_ab_tag = [i for i, t in enumerate(tags) if t in ('abstract-body', 'abstract-label')]
cn_kw_tag = [i for i, t in enumerate(tags) if t in ('keywords-body', 'keywords-label')]
en_ab_tag = [i for i, t in enumerate(tags) if t in ('en-abstract-body', 'en-abstract-label')]
en_kw_tag = [i for i, t in enumerate(tags) if t in ('en-keywords-body', 'en-keywords-label')]
has_abstract = len(cn_ab_tag) > 0
has_en_abstract = len(en_ab_tag) > 0

# ============================================================
# 第五步：自动生成目录（在摘要/关键词之后、正文第一章之前插入 TOC 域）
# ============================================================

# 找插入点：英文关键词之后 > 中文关键词之后 > 英文摘要之后 > 中文摘要之后
# TOC 必须在所有摘要（中+英）之后、正文之前
toc_insert_idx = None
all_kw = en_kw_tag + cn_kw_tag
all_ab = en_ab_tag + cn_ab_tag
if all_kw:
    toc_insert_idx = max(all_kw) + 1
elif all_ab:
    toc_insert_idx = max(all_ab) + 1

# 如果无摘要/关键词，则在第一个章节标题前插入
if toc_insert_idx is None:
    for i, t in enumerate(tags):
        if t == 'heading1' and '摘要' not in doc.paragraphs[i].text:
            toc_insert_idx = i
            break

# 防重复：扫描全文，检查是否已有 TOC 域（instrText 中包含 "TOC"）或目录标题
already_has_toc = False
for p in doc.paragraphs:
    # 检查段落 XML 中是否包含 TOC 域指令
    for instr in p._element.findall('.//' + qn('w:instrText')):
        if instr.text and 'TOC' in instr.text:
            already_has_toc = True
            break
    if already_has_toc:
        break

# 确保插入点之后没有其他 heading1 被跳过（目录插入在正文第一章前）
if toc_insert_idx is not None and not already_has_toc:
    body = doc.element.body
    ref_element = doc.paragraphs[toc_insert_idx]._element

    # ---- 目录标题段落（小三黑体，居中，2 倍行距） ----
    toc_heading_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    toc_heading_pPr = parse_xml(
        f'<w:pPr {nsdecls("w")}>'
        f'<w:jc w:val="center"/>'
        f'<w:spacing w:line="480" w:lineRule="auto" w:before="0" w:after="0"/>'
        f'</w:pPr>'
    )
    toc_heading_p.insert(0, toc_heading_pPr)
    toc_heading_r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    toc_heading_rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rFonts w:eastAsia="黑体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="30"/>'
        f'</w:rPr>'
    )
    toc_heading_r.insert(0, toc_heading_rPr)
    toc_heading_t = parse_xml(
        f'<w:t {nsdecls("w")} xml:space="preserve">目录</w:t>'
    )
    toc_heading_r.append(toc_heading_t)
    toc_heading_p.append(toc_heading_r)

    # ---- TOC 域段落 ----
    toc_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    toc_pPr = parse_xml(
        f'<w:pPr {nsdecls("w")}>'
        f'<w:spacing w:line="312" w:lineRule="auto" w:before="0" w:after="0"/>'
        f'</w:pPr>'
    )
    toc_p.insert(0, toc_pPr)

    # begin
    r_begin = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    r_begin.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    toc_p.append(r_begin)

    # instrText
    r_instr = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    r_instr.append(
        parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    )
    toc_p.append(r_instr)

    # separate
    r_sep = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    r_sep.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    toc_p.append(r_sep)

    # 提示文字
    r_hint = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    r_hint_rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rFonts w:eastAsia="宋体" w:ascii="Times New Roman"/>'
        f'<w:sz w:val="21"/>'
        f'</w:rPr>'
    )
    r_hint.insert(0, r_hint_rPr)
    r_hint.append(
        parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">（请在 Word 中右键此处 → 更新域，以生成目录）</w:t>')
    )
    toc_p.append(r_hint)

    # end
    r_end = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
    r_end.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    toc_p.append(r_end)

    # 插入到正文第一个章节之前
    ref_element.addprevious(toc_p)
    ref_element.addprevious(toc_heading_p)

    print("已自动插入目录（TOC 域），请在 Word 中更新域以生成目录。")


# ============================================================
# 第六步：表格格式化 → 五号宋体，单倍行距，居中，三线表样式
# ============================================================
# 三线表：顶线和底线为粗线（1.5pt），表头下方为细线（0.75pt），无竖线，无内横线
# OpenXML 边框单位是 1/8 pt

def remove_table_borders(table):
    """清除表格所有默认边框"""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

def set_table_border(table, position, size, color='000000'):
    """设置表格级边框（top/bottom）"""
    tblPr = table._tbl.find(qn('w:tblPr'))
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}></w:tblBorders>')
        tblPr.append(borders)
    border_el = parse_xml(
        f'<w:{position} {nsdecls("w")} '
        f'w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
    )
    borders.append(border_el)

def set_cell_bottom_border(cell, size, color='000000'):
    """给单元格添加底部边框"""
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        cell._tc.insert(0, tcPr)
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

title_page_fields = {'学院', '专业', '姓名', '学号', '指导教师',
                     '学　　院', '专　　业', '姓　　名', '学　　号'}

for idx, table in enumerate(doc.tables):
    # 判断是否为题目页表格（含学院、专业、姓名等字段）
    is_title_table = False
    if idx == 0 and len(table.rows) > 0:
        first_cell_text = ''
        if table.rows[0].cells:
            first_cell_text = table.rows[0].cells[0].text.strip()
        is_title_table = first_cell_text in title_page_fields or '学院' in first_cell_text

    if is_title_table:
        # 题目页表格：小三(15pt)宋体，无边框，单倍行距
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold or '**' in run.text:
                            set_run_font(run, '宋体', 'Times New Roman', Pt(15), True)
                        else:
                            set_run_font(run, '宋体', 'Times New Roman', Pt(15), False)
                    set_para_spacing(para, 1.5)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    clear_first_line_indent(para)
                    clear_para_spacing(para)
        remove_table_borders(table)
    else:
        # 1. 文字格式：五号宋体，单倍行距，居中
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    format_para_full(para, '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.CENTER, 1.0)

        # 2. 三线表边框
        remove_table_borders(table)
        set_table_border(table, 'top', '12')      # 1.5pt 顶线（粗）
        set_table_border(table, 'bottom', '12')   # 1.5pt 底线（粗）

        # 表头行下方加细线（0.75pt）
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                set_cell_bottom_border(cell, '6')  # 0.75pt


# ============================================================
# 第七步：摘要分节 → 中/英文摘要各成一节，实现不同页眉/页码
# ============================================================
# 规范要求：
#   中文摘要页眉 "摘要"，页码罗马数字 Ⅰ
#   英文摘要页眉 "ABSTRACT"，页码罗马数字 Ⅱ
#   正文奇数页页眉为所在章一级标题，偶数页为 "河南理工大学本科毕业设计（论文）"
#   正文页码为阿拉伯数字

# has_abstract / has_en_abstract / cn_ab_tag / cn_kw_tag / en_ab_tag / en_kw_tag 已在第四步计算

# --- 检测题目页 ---
title_page_indices = [i for i, t in enumerate(tags) if t.startswith('title-')]
has_title_page = len(title_page_indices) > 0


def insert_section_break_after(para_idx):
    """在指定段落索引处插入分节符（下一页），如已存在则跳过"""
    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        return
    para = doc.paragraphs[para_idx]
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    if pPr.find(qn('w:sectPr')) is not None:
        return
    body = doc.element.body
    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        pPr.append(copy.deepcopy(body_sectPr))
        new_sectPr = pPr.find(qn('w:sectPr'))
        if new_sectPr is not None:
            for tn in ('w:headerReference', 'w:footerReference',
                        'w:evenAndOddHeaders', 'w:titlePg'):
                for el in new_sectPr.findall(qn(tn)):
                    new_sectPr.remove(el)
        for tn in ('w:headerReference', 'w:footerReference',
                    'w:evenAndOddHeaders', 'w:titlePg'):
            for el in body_sectPr.findall(qn(tn)):
                body_sectPr.remove(el)


def insert_section_break(para):
    """在指定段落末尾插入分节符（下一页），如已存在则跳过"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para._element.insert(0, pPr)
    # 防重复：该段落已有分节符则跳过
    if pPr.find(qn('w:sectPr')) is not None:
        return
    body = doc.element.body
    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        pPr.append(copy.deepcopy(body_sectPr))
        # 清除刚复制的 sectPr 中的页眉/页脚引用，下一节将独立设置
        new_sectPr = pPr.find(qn('w:sectPr'))
        if new_sectPr is not None:
            for tn in ('w:headerReference', 'w:footerReference',
                        'w:evenAndOddHeaders', 'w:titlePg'):
                for el in new_sectPr.findall(qn(tn)):
                    new_sectPr.remove(el)
        # 旧 body 级 sectPr 也清除，准备给正文节重建
        for tn in ('w:headerReference', 'w:footerReference',
                    'w:evenAndOddHeaders', 'w:titlePg'):
            for el in body_sectPr.findall(qn(tn)):
                body_sectPr.remove(el)


# 插入分节符：题目页末尾（题目页与中文摘要/目录之间）
if has_title_page:
    last_title = max(title_page_indices)
    insert_section_break(doc.paragraphs[last_title])

# 插入分节符：中文摘要末尾
if has_abstract:
    last_cn = max(cn_ab_tag + cn_kw_tag)
    insert_section_break(doc.paragraphs[last_cn])

# 插入分节符：英文摘要末尾
if has_en_abstract:
    last_en = max(en_ab_tag + en_kw_tag)
    insert_section_break(doc.paragraphs[last_en])

# ---- 为每章插入分节符，实现奇数页页眉随章节变化 ----
# 注意：TOC 插入后段落索引已偏移，从实时 doc.paragraphs 按原始样式查找

# has_heading1 在第三步前已计算（retagging 前），用于判断原文档是否有真实 H1
# 若无 H1，则原始 H2 就是章节级标题
# 找到正文起始位置（"摘要"之后第一个章节标题），排除题目页元素
content_start = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == '摘要' or text == 'ABSTRACT' or '目录' in text:
        content_start = i + 1
        break

body_h1_indices = []
body_h1_texts = []
for i, para in enumerate(doc.paragraphs):
    if i < content_start:
        continue  # 跳过题目页和摘要区域
    text = para.text.strip()
    style = para.style.name if para.style else ''
    is_chapter = (style == 'Heading 1' or
                  (style == 'Heading 2' and not has_heading1))
    if is_chapter and '摘要' not in text and 'ABSTRACT' not in text and '目录' not in text:
        body_h1_indices.append(i)
        body_h1_texts.append(text)

chapter_breaks_added = 0
if len(body_h1_indices) >= 1:
    for h1_idx in reversed(body_h1_indices[:]):
        prev_para = doc.paragraphs[h1_idx - 1]
        pPr = prev_para._element.find(qn('w:pPr'))
        had_break = pPr is not None and pPr.find(qn('w:sectPr')) is not None
        insert_section_break(prev_para)
        if not had_break:
            chapter_breaks_added += 1
if chapter_breaks_added > 0:
    print(f"已为 {len(body_h1_indices)} 个章节插入分节符，各章奇数页页眉独立。")


# ============================================================
# 第八步：页眉和页脚设置
# ============================================================

def add_page_number(paragraph, font_name='宋体', font_size=Pt(9), en_font='Times New Roman'):
    """在段落中添加页码域（PAGE field）"""
    run_b = paragraph.add_run()
    run_b._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run_i = paragraph.add_run()
    run_i._element.append(
        parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    run_e = paragraph.add_run()
    run_e._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    for r in paragraph.runs:
        set_run_font(r, font_name, en_font, font_size)


def add_header_line(paragraph):
    """在段落底部添加页眉单横线"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        paragraph._element.insert(0, pPr)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def setup_header_footer(section, odd_header_text, even_header_text,
                         page_num_format='decimal', page_start=None):
    """
    配置一个节的页眉和页脚
    page_num_format: 'decimal' (阿拉伯数字) 或 'upperRoman' (大写罗马数字)
    page_start: 起始页码（None=沿用上一节，1=重新从1开始）
    """
    sectPr = section._sectPr

    # 开启奇偶页不同页眉
    evenAndOdd = sectPr.find(qn('w:evenAndOddHeaders'))
    if evenAndOdd is None:
        evenAndOdd = parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}></w:evenAndOddHeaders>')
        sectPr.append(evenAndOdd)

    # 设置页码格式
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")}></w:pgNumType>')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), page_num_format)
    if page_start is not None:
        pgNumType.set(qn('w:start'), str(page_start))

    # --- 偶数页页眉 ---
    even_hdr = section.even_page_header
    even_hdr.is_linked_to_previous = False
    for p in even_hdr.paragraphs:
        p.clear()
    hp_even = even_hdr.paragraphs[0] if even_hdr.paragraphs else even_hdr.add_paragraph()
    hp_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp_even.add_run(even_header_text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(9))
    add_header_line(hp_even)

    # 偶数页页脚
    even_ftr = section.even_page_footer
    even_ftr.is_linked_to_previous = False
    for p in even_ftr.paragraphs:
        p.clear()
    fp_even = even_ftr.paragraphs[0] if even_ftr.paragraphs else even_ftr.add_paragraph()
    fp_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp_even)

    # --- 奇数页页眉 ---
    odd_hdr = section.header
    odd_hdr.is_linked_to_previous = False
    for p in odd_hdr.paragraphs:
        p.clear()
    hp_odd = odd_hdr.paragraphs[0] if odd_hdr.paragraphs else odd_hdr.add_paragraph()
    hp_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp_odd.add_run(odd_header_text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(9))
    add_header_line(hp_odd)

    # 奇数页页脚
    odd_ftr = section.footer
    odd_ftr.is_linked_to_previous = False
    for p in odd_ftr.paragraphs:
        p.clear()
    fp_odd = odd_ftr.paragraphs[0] if odd_ftr.paragraphs else odd_ftr.add_paragraph()
    fp_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp_odd)


def _header_for_section(idx, title_offset, abs_offset, has_abstract, has_en_abstract):
    """根据节索引返回该节应使用的页眉文字（奇数页）。"""
    if idx < abs_offset - 1:
        # 摘要节（非最后一个前置节）
        if has_abstract and has_en_abstract:
            # 中+英：第一节=中文摘要，第二节=英文摘要
            return '摘要' if idx - title_offset == 0 else 'ABSTRACT'
        elif has_abstract:
            return '摘要'
        elif has_en_abstract:
            return 'ABSTRACT'
        else:
            return '摘要'  # fallback
    else:
        # 最后一个前置节 = 目录
        return '目录'


# 确定前置节数量（题目页 + 摘要 + 目录）
title_offset = 1 if has_title_page else 0
abs_offset = title_offset
if has_abstract:
    abs_offset += 1
if has_en_abstract:
    abs_offset += 1
abs_offset += 1  # TOC 目录节

roman_started = False
decimal_started = False

for idx, section in enumerate(doc.sections):
    # 题目页节：跳过，不设页眉页脚
    if has_title_page and idx == 0:
        continue

    is_pre_content = idx < abs_offset

    if is_pre_content:
        # 摘要 / 目录节：使用大写罗马数字页码
        hdr = _header_for_section(idx, title_offset, abs_offset,
                                  has_abstract, has_en_abstract)
        if not roman_started:
            setup_header_footer(section, hdr, hdr,
                                 page_num_format='upperRoman', page_start=1)
            roman_started = True
        else:
            setup_header_footer(section, hdr, hdr,
                                 page_num_format='upperRoman')
    else:
        # 正文节：使用阿拉伯数字页码
        chapter_idx = idx - abs_offset
        if chapter_idx < len(body_h1_texts):
            chap_title = body_h1_texts[chapter_idx]
        else:
            chap_title = odd_header_text  # fallback
        if not decimal_started:
            setup_header_footer(section, chap_title, even_header_text,
                                 page_num_format='decimal', page_start=1)
            decimal_started = True
        else:
            setup_header_footer(section, chap_title, even_header_text,
                                 page_num_format='decimal')


# ============================================================
# 第九步：保存
# ============================================================
doc.save('《地球物理综合应用与实践》报告.docx')
print("Done! 排版完成，已按河南理工大学本科毕业设计（论文）撰写规范处理。")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
if has_title_page:
    print("检测到题目页，已设置独立分节（无页眉页脚）。")
if has_abstract:
    print("检测到摘要，已插入分节符并设置罗马数字页码。")
